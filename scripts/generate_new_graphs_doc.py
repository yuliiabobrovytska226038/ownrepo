"""Generate a Word document describing the new graphs for Market Presentation 2026."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading("Nieuwe Figuren — Marktpresentatie 2026", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("Dit document beschrijft de nieuwe grafieken die zijn toegevoegd aan de Marktpresentatie 2026. Alle nieuwe figuren worden opgeslagen in de map: grafieken/nieuwe figuren/")

# ============================================================
# BATCH 1: Items 1-8
# ============================================================
doc.add_heading("Batch 1: Marktwaardeparameters & Waardedrijvers", level=1)

# Item 1
doc.add_heading("1. Verschil marktwaarde basis-full per COROP-plusgebied", level=2)
doc.add_paragraph(
    "Type: Choropleth kaart (COROP-plusgebieden)\n"
    "Bron: dataset_basis (Netto marktwaarde Basis vs Full)\n"
    "Berekening: Mediaan van (MW_basis - MW_full) / MW_full per COROP-plusgebied\n"
    "Weergave: Percentageverschil per regio, kleurschaal"
)
doc.add_paragraph("dbt model: chart_parameters__verschil_basis_full_corop.sql", style="List Bullet")
doc.add_paragraph("Bestandsnaam: Verschil marktwaarde basis-full per COROP-plusgebied.html", style="List Bullet")

# Item 2 (same as item 1, part of chart_parameters)
# Actually item 2 was already in the parameters charts

# Item 3
doc.add_heading("2. Waardedrijvers wolk — cumulatief effect (Marktwaarde)", level=2)
doc.add_paragraph(
    "Type: Bubble chart (scatter met grootte)\n"
    "Bron: chart_va__marketvalue_drivers (verschillenanalyse)\n"
    "X-as: Cumulatief effect per stap\n"
    "Y-as: Stap (waardedrijver)\n"
    "Grootte: Absolute mutatie per stap\n"
    "Kleur: Waarderingstype (Basis/Full)"
)
doc.add_paragraph("Bestandsnaam: Marktwaardedrijvers wolk cumulatief effect.html", style="List Bullet")

# Item 3b
doc.add_heading("3. Waardedrijvers wolk — cumulatief effect (Beleidswaarde)", level=2)
doc.add_paragraph(
    "Type: Bubble chart (scatter met grootte)\n"
    "Bron: chart_va__policyvalue_drivers (verschillenanalyse)\n"
    "X-as: Cumulatief effect per stap\n"
    "Y-as: Stap (waardedrijver)\n"
    "Grootte: Absolute mutatie per stap\n"
    "Kleur: Waarderingstype"
)
doc.add_paragraph("Bestandsnaam: Beleidswaardedrijvers wolk cumulatief effect.html", style="List Bullet")

# Item 4
doc.add_heading("4. Percentage uitponden per COROP-plusgebied (Basis & Full)", level=2)
doc.add_paragraph(
    "Type: 2x Choropleth kaart (COROP-plusgebieden)\n"
    "Bron: dataset_basis (Scenario uitponden/totaal per COROP)\n"
    "Berekening: % VHE met scenario uitponden per COROP, apart voor Basis en Full\n"
    "Weergave: Percentages per regio"
)
doc.add_paragraph("dbt model: chart_parameters__uitponden_corop.sql", style="List Bullet")
doc.add_paragraph("Bestandsnamen: Percentage uitponden per COROP-plusgebied Basis.html, ...Full.html", style="List Bullet")

# Item 5
doc.add_heading("5. Ontwikkeling percentage uitponden per COROP-plusgebied (Basis & Full)", level=2)
doc.add_paragraph(
    "Type: 2x Choropleth kaart (COROP-plusgebieden)\n"
    "Bron: dataset_basis (huidig jaar vs vorig jaar)\n"
    "Berekening: Verschil in uitpondpercentage t.o.v. vorig jaar per COROP\n"
    "Weergave: Mutatie in procentpunten per regio"
)
doc.add_paragraph("dbt model: chart_parameters__uitponden_ontwikkeling_corop.sql", style="List Bullet")
doc.add_paragraph("Bestandsnamen: Ontwikkeling percentage uitponden per COROP-plusgebied Basis.html, ...Full.html", style="List Bullet")

# Item 6
doc.add_heading("6. Ontwikkeling disconteringsvoet per COROP-plusgebied (Basis, Full & Totaal)", level=2)
doc.add_paragraph(
    "Type: 3x Choropleth kaart (COROP-plusgebieden)\n"
    "Bron: dataset_ontwikkeling (DV doorexploiteren huidig vs vorig jaar)\n"
    "Berekening: Mediaan DV huidig - mediaan DV vorig jaar per COROP\n"
    "Weergave: Mutatie in procentpunten per regio\n"
    "Varianten: Basis, Full, en Totaal (alle waarderingstypen gecombineerd)"
)
doc.add_paragraph("dbt model: chart_parameters__dv_ontwikkeling_corop.sql", style="List Bullet")
doc.add_paragraph("Bestandsnamen: Ontwikkeling disconteringsvoet per COROP-plusgebied Basis.html, ...Full.html, ...Totaal.html", style="List Bullet")

# Item 7
doc.add_heading("7. Gehanteerde disconteringsvoet per COROP-plusgebied (Basis, Full & Totaal)", level=2)
doc.add_paragraph(
    "Type: 3x Choropleth kaart (COROP-plusgebieden)\n"
    "Bron: dataset_basis (DV doorexploiteren per COROP)\n"
    "Berekening: Mediaan DV doorexploiteren per COROP, apart voor Basis, Full en Totaal\n"
    "Weergave: Percentage per regio\n"
    "Varianten: Basis, Full, en Totaal (alle waarderingstypen gecombineerd — altijd 52/52 regio's ingekleurd)"
)
doc.add_paragraph("dbt model: chart_parameters__dv_corop.sql", style="List Bullet")
doc.add_paragraph("Bestandsnamen: Gehanteerde disconteringsvoet per COROP-plusgebied Basis.html, ...Full.html, ...Totaal.html", style="List Bullet")

# Item 8
doc.add_heading("8. Trendlijnen parameters per waarderingstype", level=2)
doc.add_paragraph(
    "Type: Lijngrafieken (4 parameters × 2 varianten = 8 grafieken)\n"
    "Bron: dataset_basis + dataset_ontwikkeling (historische jaren)\n"
    "Parameters: Leegwaarde, Markthuur, Exit yield, Disconteringsvoet\n"
    "Varianten per parameter:\n"
    "  a) Trendlijn per waarderingstype (Basis vs Full over de jaren)\n"
    "  b) Trendlijn verschil basis-full over de jaren"
)
doc.add_paragraph("dbt model: chart_parameters__trendlijnen.sql", style="List Bullet")
doc.add_paragraph("Bestandsnamen: Trendlijn [Parameter] per waarderingstype.html, Trendlijn verschil basis-full [Parameter].html", style="List Bullet")

# ============================================================
# BATCH 2: Items 9-18
# ============================================================
doc.add_heading("Batch 2: Beleidswaardeparameters", level=1)

# Item 9
doc.add_heading("9. Waardedrijvers beleidswaarde — wolk cumulatief effect", level=2)
doc.add_paragraph(
    "Type: Bubble chart (scatter met grootte)\n"
    "Bron: chart_va__policyvalue_drivers (TMS verschillenanalyse beleidswaarde)\n"
    "Stappen in de data: Onderhoudskosten beleidswaarde, Beheerkosten beleidswaarde, "
    "Beleidshuurbeleid, Contracthuur, Disconteringsvoet beleidswaarde, Overige mutaties\n"
    "X-as: Cumulatief effect\n"
    "Y-as: Stap (waardedrijver)\n"
    "Grootte: Absolute mutatie\n"
    "Opmerking: Reeds geïmplementeerd in batch 1 (chart_va.py)"
)
doc.add_paragraph("Bestandsnaam: Beleidswaardedrijvers wolk cumulatief effect.html", style="List Bullet")

# Item 10
doc.add_heading("10. Kaart ontwikkeling beleidsonderhoud per VHE per COROP-plusgebied", level=2)
doc.add_paragraph(
    "Type: Choropleth kaart (COROP-plusgebieden)\n"
    "Bron: dataset_ontwikkeling (Beleidsonderhoud huidig vs vorig jaar)\n"
    "Berekening: Mediaan beleidsonderhoud huidig - mediaan vorig jaar per COROP\n"
    "Weergave: Verschil in euro's per maand per regio\n"
    "Filter: Alleen Woningen, _merge = 'both'"
)
doc.add_paragraph("dbt model: chart_blw_params__beleidsonderhoud_ontwikkeling_corop.sql", style="List Bullet")
doc.add_paragraph("Bestandsnaam: Ontwikkeling beleidsonderhoud per COROP-plusgebied.html", style="List Bullet")

# Item 11
doc.add_heading("11. Kaart ontwikkeling beleidsbeheer per VHE per COROP-plusgebied", level=2)
doc.add_paragraph(
    "Type: Choropleth kaart (COROP-plusgebieden)\n"
    "Bron: dataset_ontwikkeling (Beleidsbeheer_bp huidig vs Beleidsbeheer_2024)\n"
    "Berekening: Mediaan beleidsbeheer huidig - mediaan vorig jaar per COROP\n"
    "Weergave: Verschil in euro's per maand per regio\n"
    "Filter: Alleen Woningen, _merge = 'both'"
)
doc.add_paragraph("dbt model: chart_blw_params__beleidsbeheer_ontwikkeling_corop.sql", style="List Bullet")
doc.add_paragraph("Bestandsnaam: Ontwikkeling beleidsbeheer per COROP-plusgebied.html", style="List Bullet")

# Item 12
doc.add_heading("12. Kaart ontwikkeling beleidshuur per VHE per COROP-plusgebied", level=2)
doc.add_paragraph(
    "Type: Choropleth kaart (COROP-plusgebieden)\n"
    "Bron: dataset_ontwikkeling (Beleidshuur huidig vs Beleidshuur_2024)\n"
    "Berekening: Mediaan beleidshuur huidig - mediaan vorig jaar per COROP\n"
    "Weergave: Verschil in euro's per maand per regio\n"
    "Filter: Alleen Woningen, _merge = 'both'"
)
doc.add_paragraph("dbt model: chart_blw_params__beleidshuur_ontwikkeling_corop.sql", style="List Bullet")
doc.add_paragraph("Bestandsnaam: Ontwikkeling beleidshuur per COROP-plusgebied.html", style="List Bullet")

# Item 13
doc.add_heading("13. Verdeling beleidshuren over huurcategorieën", level=2)
doc.add_paragraph(
    "Type: Taartdiagram (pie chart)\n"
    "Bron: tms.policy_value_parameters (Segment huurregime)\n"
    "Categorieën: Laagsegment, Middensegment, Hoogsegment\n"
    "Berekening: Aantal VHE per segment + percentage\n"
    "Filter: Huidig jaar (2025)"
)
doc.add_paragraph("dbt model: chart_blw_params__verdeling_huurcategorieen.sql", style="List Bullet")
doc.add_paragraph("Bestandsnaam: Verdeling beleidshuren over huurcategorieen.html", style="List Bullet")

# Item 14
doc.add_heading("14. Ontwikkeling verdeling beleidshuren over huurcategorieën", level=2)
doc.add_paragraph(
    "Type: Gegroepeerd staafdiagram (grouped bar chart)\n"
    "Bron: tms.policy_value_parameters (Segment huurregime, 2024 vs 2025)\n"
    "X-as: Segment huurregime\n"
    "Y-as: Percentage (aandeel)\n"
    "Groepen: Huidig jaar vs vorig jaar\n"
    "Berekening: Percentage VHE per segment per jaar, verschil"
)
doc.add_paragraph("dbt model: chart_blw_params__verdeling_huurcategorieen_ontwikkeling.sql", style="List Bullet")
doc.add_paragraph("Bestandsnaam: Ontwikkeling verdeling beleidshuren over huurcategorieen.html", style="List Bullet")

# Item 15
doc.add_heading("15. Lengte MJOB", level=2)
doc.add_paragraph(
    "Type: Staafdiagram (bar chart)\n"
    "Bron: tms.policy_value_parameters (Beleidsonderhoud jaar 1 t/m 60)\n"
    "X-as: Lengte MJOB groep (gegroepeerd in categorieën)\n"
    "Y-as: Aantal VHE\n"
    "Berekening: Per VHE wordt geteld hoeveel van de 60 jaarvelden een waarde > 0 hebben, "
    "vervolgens gegroepeerd in bins: 0 (geen), 1-15, 16-30, 31-45, 46-59, 60 (volledig)\n"
    "Filter: Huidig jaar (2025)\n"
    "Opmerking: 83% van de VHEs heeft volledige MJOB (60 jaar), ~6% heeft geen MJOB"
)
doc.add_paragraph("dbt model: chart_blw_params__lengte_mjob.sql", style="List Bullet")
doc.add_paragraph("Bestandsnaam: Lengte MJOB verdeling.html", style="List Bullet")

# Item 16
doc.add_heading("16. Boxplot realisatie reguliere huurstijging", level=2)
doc.add_paragraph(
    "Type: Horizontale boxplot\n"
    "Bron: tms.vastgoedgegevens_vhe_gegevens + tms.policy_value_parameters\n"
    "X-as: Verschil realisatie - verwachte huurstijging (procentpunten)\n"
    "Eenheid per datapunt: Corporatie\n"
    "Berekening per corporatie:\n"
    "  1. Gemeten huurstijging = (Netto huur 2025 - Netto huur 2024) / Netto huur 2024\n"
    "  2. Verwachte huurstijging = HS jaar 1 uit beleidswaardeparameters 2024\n"
    "  3. Verschil = gemeten - verwacht\n"
    "  4. Mediaan van het verschil over alle VHEs van die corporatie\n"
    "Filter: Alleen EGW/MGW met gemeten huurstijging tussen 0% en 5%\n"
    "Annotatie: Gemiddelde als rode stippellijn"
)
doc.add_paragraph("dbt model: chart_blw_params__realisatie_huurstijging.sql", style="List Bullet")
doc.add_paragraph("Bestandsnaam: Boxplot realisatie reguliere huurstijging.html", style="List Bullet")

# Item 17
doc.add_heading("17. Boxplot realisatie beleidshuur", level=2)
doc.add_paragraph(
    "Type: Horizontale boxplot\n"
    "Bron: tms.vastgoedgegevens_vhe_gegevens + tms.policy_value_parameters\n"
    "X-as: % woningen waar netto huur = beleidshuur\n"
    "Eenheid per datapunt: Corporatie\n"
    "Berekening per corporatie:\n"
    "  1. Filter op mutaties: VHEs met gemeten huurstijging < 0% of > 5%\n"
    "  2. Per mutatie-VHE: check of netto huur (VGG 2025) = beleidshuur (BLW params 2025)\n"
    "  3. Percentage = # gelijk / # mutatie-VHEs\n"
    "Filter: Alleen EGW/MGW\n"
    "Annotatie: Gemiddelde als rode stippellijn"
)
doc.add_paragraph("dbt model: chart_blw_params__realisatie_beleidshuur.sql", style="List Bullet")
doc.add_paragraph("Bestandsnaam: Boxplot realisatie beleidshuur.html", style="List Bullet")

# Item 18
doc.add_heading("18. Boxplot realisatie MJOB", level=2)
doc.add_paragraph(
    "Type: Horizontale boxplot\n"
    "Bron: tms.vastgoedgegevens_vhe_gegevens + tms.policy_value_parameters\n"
    "X-as: Verschil realisatie - verwacht beleidsonderhoud (procentpunten)\n"
    "Eenheid per datapunt: Corporatie\n"
    "Berekening per corporatie:\n"
    "  1. Verwacht onderhoud = Beleidsonderhoud jaar 1 (BLW params 2024)\n"
    "  2. Actueel onderhoud = Beleidsonderhoud (dataset_ontwikkeling, huidig jaar)\n"
    "  3. Verschil = (actueel - verwacht) / verwacht\n"
    "  4. Mediaan verschil over alle VHEs\n"
    "Filter: Alleen Woningen met verwacht onderhoud > 0\n"
    "Annotatie: Gemiddelde als rode stippellijn"
)
doc.add_paragraph("dbt model: chart_blw_params__realisatie_mjob.sql", style="List Bullet")
doc.add_paragraph("Bestandsnaam: Boxplot realisatie MJOB.html", style="List Bullet")

# ============================================================
# BATCH 3: Conclusie dashboards
# ============================================================
doc.add_heading("Batch 3: Conclusie dashboards", level=1)

doc.add_paragraph(
    "Visuele samenvattingen van de belangrijkste conclusies per waarderingsonderdeel. Geen lap tekst maar KPI-scorecards met indicatoren, kleurcodes en watervalgrafiek van de waardedrijvers."
)

doc.add_heading("19. Conclusie Marktwaarde dashboard", level=2)
doc.add_paragraph(
    "Type: Dashboard (Plotly subplots met Indicator + Bar)\n"
    "Opbouw:\n"
    "  Rij 1 — KPI-tegels:\n"
    "    • Waardestijging Basis (% met delta-indicator)\n"
    "    • Waardestijging Full (% met delta-indicator)\n"
    "    • Mediaan disconteringsvoet (%)\n"
    "    • Mediaan exit yield (%)\n"
    "  Rij 2 — Waardedrijvers waterval:\n"
    "    • Horizontale balken per stap, gesorteerd op impact\n"
    "    • Groen = positieve bijdrage, Rood = negatieve bijdrage\n"
    "Bron: dataset_ontwikkeling, dataset_basis, chart_va__marketvalue_drivers\n"
    "Conclusie in één oogopslag: welke parameter dreef de waarde omhoog/omlaag"
)
doc.add_paragraph("dbt model: chart_conclusie__marktwaarde.sql", style="List Bullet")
doc.add_paragraph("Bestandsnaam: Conclusie Marktwaarde dashboard.html", style="List Bullet")

doc.add_heading("20. Conclusie Beleidswaarde dashboard", level=2)
doc.add_paragraph(
    "Type: Dashboard (Plotly subplots met Indicator + Bar)\n"
    "Opbouw:\n"
    "  Rij 1 — KPI-tegels:\n"
    "    • Waardestijging BLW (% met delta-indicator)\n"
    "    • Ratio BLW/MW (mediaan %)\n"
    "    • Beleidsonderhoud €/mnd (met delta t.o.v. vorig jaar)\n"
    "    • Beleidshuur €/mnd (met delta t.o.v. vorig jaar)\n"
    "  Rij 2 — Waardedrijvers waterval:\n"
    "    • Horizontale balken per BLW-driver, gesorteerd op impact\n"
    "    • Groen = positief, Rood = negatief\n"
    "  Rij 3 — Parameterontwikkeling:\n"
    "    • % mutatie Beleidsonderhoud, Beleidsbeheer, Beleidshuur t.o.v. vorig jaar\n"
    "    • Kleurgecodeerd (groen/rood)\n"
    "Bron: dataset_ontwikkeling, dataset_basis, chart_va__policyvalue_drivers, policy_value_parameters\n"
    "Conclusie in één oogopslag: BLW-ontwikkeling + welke parameters het meest verschoven"
)
doc.add_paragraph("dbt model: chart_conclusie__beleidswaarde.sql", style="List Bullet")
doc.add_paragraph("Bestandsnaam: Conclusie Beleidswaarde dashboard.html", style="List Bullet")

# ============================================================
# Technical Summary
# ============================================================
doc.add_heading("Technische samenvatting", level=1)

doc.add_heading("Dagster assets", level=2)
table = doc.add_table(rows=5, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Asset"
hdr[1].text = "Module"
hdr[2].text = "Items"
rows_data = [
    ("charts_parameters", "chart_parameters.py", "1, 4-8"),
    ("charts_va", "chart_va.py", "2-3, 9"),
    ("charts_blw_params", "chart_blw_params.py", "10-18"),
    ("charts_conclusies", "chart_conclusies.py", "19-20"),
]
for i, (asset, module, items) in enumerate(rows_data, 1):
    row = table.rows[i].cells
    row[0].text = asset
    row[1].text = module
    row[2].text = items

doc.add_paragraph("")

doc.add_heading("dbt modellen", level=2)
doc.add_paragraph("Batch 1 (chart_parameters__*):", style="List Bullet")
doc.add_paragraph("verschil_basis_full_corop, uitponden_corop, uitponden_ontwikkeling_corop, dv_corop, dv_ontwikkeling_corop, trendlijnen", style="List Bullet 2")
doc.add_paragraph("Batch 2 (chart_blw_params__*):", style="List Bullet")
doc.add_paragraph(
    "beleidsonderhoud_ontwikkeling_corop, beleidsbeheer_ontwikkeling_corop, beleidshuur_ontwikkeling_corop, verdeling_huurcategorieen, verdeling_huurcategorieen_ontwikkeling, lengte_mjob, realisatie_huurstijging, realisatie_beleidshuur, realisatie_mjob",
    style="List Bullet 2",
)
doc.add_paragraph("Batch 3 (chart_conclusie__*):", style="List Bullet")
doc.add_paragraph("marktwaarde, beleidswaarde", style="List Bullet 2")

doc.add_heading("Materialisatie", level=2)
doc.add_paragraph(
    "1. Materialiseer TMS assets (tms_assets_job) voor ruwe data\n"
    "2. Run dbt transformaties: uv run dg launch --job dbt_transformations_job\n"
    "3. Genereer grafieken: uv run dg launch --job charts_job\n"
    "4. Output staat in: grafieken/nieuwe figuren/"
)

doc.add_heading("Bugfixes COROP kaarten (mei 2026)", level=2)
doc.add_paragraph("De volgende problemen met grijze (lege) regio's op COROP-kaarten zijn opgelost:")
doc.add_paragraph(
    "corop_medians model: Filterde voorheen alleen op 'Full' waarderingstype voor mediaan LW/DV/BO. "
    "Hierdoor was CP4003 (Noordoostpolder en Urk) grijs — deze regio heeft uitsluitend 'Basis' VHEs. "
    "Fix: Alle waarderingstypen worden nu meegenomen (beide typen hebben valide LW/DV/BO waarden).",
    style="List Bullet",
)
doc.add_paragraph(
    "EFG COROP model: CP4001 (Almere) was hardcoded uitgesloten met een NOT IN filter. Fix: Uitsluiting verwijderd — nu 52/52 regio's.",
    style="List Bullet",
)
doc.add_paragraph(
    "DV COROP kaarten: 8 regio's (o.a. Amsterdam, Zuid-Limburg) waren grijs op de 'Basis' kaart "
    "omdat deze regio's uitsluitend 'Full' waarderingen bevatten, en vice versa voor CP4003 op de 'Full' kaart. "
    "Fix: Een 'Totaal' variant toegevoegd die alle waarderingstypen combineert — altijd volledige dekking.",
    style="List Bullet",
)
doc.add_paragraph(
    "Excel detail vs kaart inconsistentie: Detail Excel toonde soms data voor regio's die grijs waren op de kaart "
    "(door .dropna() vóór multi_layer_map maar volledige data naar _detail_for). "
    "Fix: Detail data nu consistent gefilterd met dezelfde subset als de kaart.",
    style="List Bullet",
)

doc.save(r"C:\Dev\rev-market-presentation\docs\Nieuwe figuren marktpresentatie 2026.docx")
print("Document saved successfully.")
